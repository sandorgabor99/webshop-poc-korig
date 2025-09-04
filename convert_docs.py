#!/usr/bin/env python3
"""
Documentation Converter Script
Converts Markdown files to Word (.docx) and PDF formats
"""

import os
import sys
from pathlib import Path
import markdown2
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import webbrowser
import tempfile

def markdown_to_html(markdown_content):
    """Convert Markdown to HTML"""
    return markdown2.markdown(markdown_content, extras=['tables', 'fenced-code-blocks', 'code-friendly'])

def html_to_word(html_content, output_file):
    """Convert HTML to Word document"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('WebShop POC - Technical Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content
    doc.add_paragraph(html_content)
    
    # Save document
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")

def html_to_pdf_via_browser(html_content, output_file):
    """Convert HTML to PDF using browser print functionality"""
    # Create a temporary HTML file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
        html_template = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>WebShop POC - Technical Documentation</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        h3 {{ color: #7f8c8d; }}
        code {{ background-color: #f8f9fa; padding: 2px 4px; border-radius: 3px; font-family: 'Courier New', monospace; }}
        pre {{ background-color: #f8f9fa; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
        .highlight {{ background-color: #fff3cd; padding: 15px; border-radius: 5px; border-left: 4px solid #ffc107; }}
    </style>
</head>
<body>
    {html_content}
    <script>
        // Auto-print when page loads
        window.onload = function() {{
            window.print();
        }};
    </script>
</body>
</html>
        """
        f.write(html_template)
        temp_file = f.name
    
    # Open in browser for printing to PDF
    print(f"🌐 Opening {temp_file} in browser...")
    print("📋 Instructions:")
    print("   1. Use Ctrl+P (or Cmd+P on Mac) to print")
    print("   2. Select 'Save as PDF' as destination")
    print("   3. Save as: {output_file}")
    print("   4. Close the browser tab when done")
    
    webbrowser.open(f'file://{temp_file}')
    
    # Clean up temp file after a delay
    import time
    time.sleep(2)
    try:
        os.unlink(temp_file)
    except:
        pass

def convert_markdown_file(markdown_file, output_dir):
    """Convert a single Markdown file to Word and PDF"""
    print(f"\n🔄 Converting: {markdown_file}")
    
    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Convert to HTML
    html_content = markdown_to_html(markdown_content)
    
    # Generate output filenames
    base_name = Path(markdown_file).stem
    word_file = output_dir / f"{base_name}.docx"
    pdf_file = output_dir / f"{base_name}.pdf"
    
    # Convert to Word
    try:
        html_to_word(html_content, word_file)
    except ImportError:
        print("⚠️  python-docx not installed. Installing...")
        os.system("pip install python-docx")
        try:
            html_to_word(html_content, word_file)
        except Exception as e:
            print(f"❌ Error creating Word document: {e}")
    
    # Convert to PDF
    try:
        html_to_pdf_via_browser(html_content, pdf_file)
    except Exception as e:
        print(f"❌ Error creating PDF: {e}")

def main():
    """Main conversion function"""
    print("🚀 WebShop POC Documentation Converter")
    print("=" * 50)
    
    # Create output directory
    output_dir = Path("converted_docs")
    output_dir.mkdir(exist_ok=True)
    
    # Find all markdown files in docs directory
    docs_dir = Path("docs")
    if not docs_dir.exists():
        print("❌ docs directory not found!")
        return
    
    markdown_files = list(docs_dir.glob("*.md"))
    
    if not markdown_files:
        print("❌ No markdown files found in docs directory!")
        return
    
    print(f"📁 Found {len(markdown_files)} markdown files:")
    for file in markdown_files:
        print(f"   - {file.name}")
    
    print(f"\n📂 Output directory: {output_dir.absolute()}")
    
    # Convert each file
    for markdown_file in markdown_files:
        convert_markdown_file(markdown_file, output_dir)
    
    print(f"\n🎉 Conversion complete! Check the '{output_dir}' folder for your documents.")
    print("\n📋 Next steps:")
    print("   1. Open the HTML files in your browser")
    print("   2. Use Ctrl+P to print and save as PDF")
    print("   3. Word documents (.docx) are ready to use")

if __name__ == "__main__":
    main()
