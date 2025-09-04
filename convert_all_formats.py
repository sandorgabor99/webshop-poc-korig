#!/usr/bin/env python3
"""
Comprehensive Documentation Converter
Converts Markdown files to Word (.docx), PDF, and HTML formats
"""

import os
import sys
from pathlib import Path
import markdown2
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import webbrowser
import tempfile
import subprocess
import time

class DocConverter:
    def __init__(self):
        self.output_dir = Path("converted_docs")
        self.output_dir.mkdir(exist_ok=True)
        
        # Create subdirectories for different formats
        self.word_dir = self.output_dir / "word"
        self.pdf_dir = self.output_dir / "pdf"
        self.html_dir = self.output_dir / "html"
        
        for dir_path in [self.word_dir, self.pdf_dir, self.html_dir]:
            dir_path.mkdir(exist_ok=True)
    
    def markdown_to_html(self, markdown_content):
        """Convert Markdown to HTML with enhanced features"""
        return markdown2.markdown(
            markdown_content, 
            extras=[
                'tables', 
                'fenced-code-blocks', 
                'code-friendly',
                'cuddled-lists',
                'markdown-in-html',
                'task_list'
            ]
        )
    
    def create_styled_html(self, html_content, title):
        """Create a styled HTML document"""
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f8f9fa;
        }}
        
        .container {{
            background: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 15px;
            margin-bottom: 30px;
            font-size: 2.5em;
            text-align: center;
        }}
        
        h2 {{
            color: #34495e;
            margin-top: 40px;
            margin-bottom: 20px;
            font-size: 1.8em;
            border-left: 4px solid #3498db;
            padding-left: 15px;
        }}
        
        h3 {{
            color: #7f8c8d;
            margin-top: 25px;
            margin-bottom: 15px;
            font-size: 1.4em;
        }}
        
        h4 {{
            color: #95a5a6;
            margin-top: 20px;
            margin-bottom: 10px;
            font-size: 1.2em;
        }}
        
        p {{
            margin-bottom: 15px;
            text-align: justify;
        }}
        
        code {{
            background-color: #f8f9fa;
            padding: 3px 6px;
            border-radius: 4px;
            font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
            font-size: 0.9em;
            border: 1px solid #e9ecef;
        }}
        
        pre {{
            background-color: #f8f9fa;
            padding: 20px;
            border-radius: 6px;
            overflow-x: auto;
            border: 1px solid #e9ecef;
            margin: 20px 0;
        }}
        
        pre code {{
            background: none;
            padding: 0;
            border: none;
            font-size: 0.95em;
        }}
        
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
            background: white;
            border-radius: 6px;
            overflow: hidden;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }}
        
        th, td {{
            border: 1px solid #dee2e6;
            padding: 12px;
            text-align: left;
        }}
        
        th {{
            background-color: #3498db;
            color: white;
            font-weight: 600;
        }}
        
        tr:nth-child(even) {{
            background-color: #f8f9fa;
        }}
        
        tr:hover {{
            background-color: #e9ecef;
        }}
        
        ul, ol {{
            margin: 15px 0;
            padding-left: 30px;
        }}
        
        li {{
            margin-bottom: 8px;
        }}
        
        blockquote {{
            border-left: 4px solid #3498db;
            padding-left: 20px;
            margin: 20px 0;
            font-style: italic;
            color: #6c757d;
        }}
        
        .highlight {{
            background-color: #fff3cd;
            padding: 15px;
            border-radius: 6px;
            border-left: 4px solid #ffc107;
            margin: 20px 0;
        }}
        
        .info {{
            background-color: #d1ecf1;
            padding: 15px;
            border-radius: 6px;
            border-left: 4px solid #17a2b8;
            margin: 20px 0;
        }}
        
        .warning {{
            background-color: #f8d7da;
            padding: 15px;
            border-radius: 6px;
            border-left: 4px solid #dc3545;
            margin: 20px 0;
        }}
        
        @media print {{
            body {{ background: white; }}
            .container {{ box-shadow: none; padding: 20px; }}
            h1, h2, h3 {{ page-break-after: avoid; }}
            pre, table {{ page-break-inside: avoid; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        {html_content}
    </div>
    
    <script>
        // Add print functionality
        function printDocument() {{
            window.print();
        }}
        
        // Add keyboard shortcut (Ctrl+P)
        document.addEventListener('keydown', function(e) {{
            if (e.ctrlKey && e.key === 'p') {{
                e.preventDefault();
                printDocument();
            }}
        }});
        
        // Auto-print after 2 seconds (optional)
        // setTimeout(printDocument, 2000);
    </script>
</body>
</html>"""
    
    def html_to_word(self, html_content, output_file):
        """Convert HTML to Word document with better formatting"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('WebShop POC - Technical Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content (simplified HTML to text conversion)
        # Remove HTML tags for Word document
        import re
        clean_text = re.sub(r'<[^>]+>', '', html_content)
        clean_text = re.sub(r'&[^;]+;', '', clean_text)
        
        # Split into paragraphs and add to document
        paragraphs = clean_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save document
        doc.save(output_file)
        print(f"✅ Word document created: {output_file}")
    
    def html_to_pdf_browser(self, html_content, output_file):
        """Convert HTML to PDF using browser print functionality"""
        # Create a temporary HTML file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
            f.write(html_content)
            temp_file = f.name
        
        print(f"🌐 Opening HTML in browser for PDF conversion...")
        print(f"📋 Instructions:")
        print(f"   1. Use Ctrl+P to print")
        print(f"   2. Select 'Save as PDF' as destination")
        print(f"   3. Save as: {output_file}")
        print(f"   4. Close the browser tab when done")
        
        # Open in browser
        webbrowser.open(f'file://{temp_file}')
        
        # Clean up temp file after delay
        time.sleep(3)
        try:
            os.unlink(temp_file)
        except:
            pass
    
    def try_wkhtmltopdf(self, html_file, pdf_file):
        """Try to use wkhtmltopdf if available"""
        # Try multiple possible paths
        wkhtmltopdf_paths = [
            'wkhtmltopdf',  # If in PATH
            r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe',  # Windows default
            r'C:\Program Files (x86)\wkhtmltopdf\bin\wkhtmltopdf.exe'  # 32-bit
        ]
        
        for wkhtmltopdf_path in wkhtmltopdf_paths:
            try:
                result = subprocess.run([
                    wkhtmltopdf_path,
                    '--page-size', 'A4',
                    '--margin-top', '20',
                    '--margin-right', '20',
                    '--margin-bottom', '20',
                    '--margin-left', '20',
                    '--encoding', 'UTF-8',
                    html_file,
                    pdf_file
                ], capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0:
                    print(f"✅ PDF created with wkhtmltopdf: {pdf_file}")
                    return True
                else:
                    print(f"⚠️  wkhtmltopdf failed: {result.stderr}")
                    continue  # Try next path
            except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError):
                continue  # Try next path
        
        return False  # All paths failed
    
    def convert_file(self, markdown_file):
        """Convert a single Markdown file to all formats"""
        print(f"\n🔄 Converting: {markdown_file}")
        
        # Read markdown content
        with open(markdown_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Convert to HTML
        html_content = self.markdown_to_html(markdown_content)
        
        # Generate output filenames
        base_name = Path(markdown_file).stem
        word_file = self.word_dir / f"{base_name}.docx"
        pdf_file = self.pdf_dir / f"{base_name}.pdf"
        html_file = self.html_dir / f"{base_name}.html"
        
        # Create styled HTML
        styled_html = self.create_styled_html(html_content, base_name)
        
        # Save HTML file
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(styled_html)
        print(f"✅ HTML file created: {html_file}")
        
        # Convert to Word
        try:
            self.html_to_word(html_content, word_file)
        except Exception as e:
            print(f"❌ Error creating Word document: {e}")
        
        # Try to create PDF
        if not self.try_wkhtmltopdf(html_file, pdf_file):
            # Fallback to browser method
            self.html_to_pdf_browser(styled_html, pdf_file)
    
    def convert_all(self):
        """Convert all markdown files"""
        print("🚀 WebShop POC - Comprehensive Documentation Converter")
        print("=" * 60)
        
        # Find all markdown files in docs directory
        docs_dir = Path("docs")
        if not docs_dir.exists():
            print("❌ docs directory not found!")
            return
        
        markdown_files = list(docs_dir.glob("*.md"))
        
        if not markdown_files:
            print("❌ No markdown files found in docs directory!")
            return
        
        print(f"📁 Found {len(markdown_files)} markdown files:")
        for file in markdown_files:
            print(f"   - {file.name}")
        
        print(f"\n📂 Output directories:")
        print(f"   Word documents: {self.word_dir.absolute()}")
        print(f"   PDF documents: {self.pdf_dir.absolute()}")
        print(f"   HTML files: {self.html_dir.absolute()}")
        
        # Convert each file
        for markdown_file in markdown_files:
            self.convert_file(markdown_file)
        
        print(f"\n🎉 Conversion complete!")
        print(f"\n📋 Summary:")
        print(f"   Word documents: {len(list(self.word_dir.glob('*.docx')))}")
        print(f"   HTML files: {len(list(self.html_dir.glob('*.html')))}")
        print(f"   PDF files: {len(list(self.pdf_dir.glob('*.pdf')))}")
        
        print(f"\n📋 Next steps:")
        print(f"   1. Word documents (.docx) are ready to use")
        print(f"   2. HTML files can be opened in browser")
        print(f"   3. For PDFs: Open HTML files and use Ctrl+P → Save as PDF")
        print(f"   4. Or install wkhtmltopdf for automatic PDF generation")

def main():
    """Main function"""
    converter = DocConverter()
    converter.convert_all()

if __name__ == "__main__":
    main()
